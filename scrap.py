import urllib3
from docx import Document
from bs4 import BeautifulSoup

# thk h hogya :D haan nice isko github pe daalenge okay? :D tu hi daalio apne me -_- saath mein kiya na
# thappad merko apni gmail ka pass bataeyo hackathons mein bhi register karna tha batade :P27316812

codechefURLS = [
    'https://www.codechef.com/problems/school',
    'https://www.codechef.com/problems/easy',
    'https://www.codechef.com/problems/medium',
    'https://www.codechef.com/problems/hard',
    'https://www.codechef.com/problems/challenge',
    'https://www.codechef.com/problems/extcontest'
]

file_headings = [
    'Codechef Beginner Problems',
    'Codechef Easy Problems',
    'Codechef Medium Problems',
    'Codechef Hard Problems',
    'Codechef Challenge Problems',
    'Codechef Peer Problems',
]

file_names = [
    'codechef_beginner',
    'codechef_easy',
    'codechef_medium',
    'codechef_hard',
    'codechef_challenge',
    'codechef_peer'
]

def PrintProblemsFromURL(url, file_heading, file_name):
    # init http object
    http = urllib3.PoolManager()

    # fetching data from url
    response = http.request('GET', url)
    html = response.data

    # init soup object
    soup = BeautifulSoup(html, 'html.parser')
    problems = soup.find_all('tr', attrs={'class': 'problemrow'})

    # lists init
    heading = []
    code = []
    successful_submissions = []
    accuracy = []

    # scrap data and append to lists
    for i in range(len(problems)):
        cur_problem = problems[i]
        table_data = cur_problem.find_all('td')
        heading.append(table_data[0].div.a.b.string)
        if table_data[1].a == None:
            code.append(table_data[1].string)
        else:
            code.append(table_data[1].a.string)
        successful_submissions.append(table_data[2].div.string)
        accuracy.append(table_data[3].a.string)

    # for i in range(len(problems)):
    #     print('heading: ' + heading[i])
    #     print('code: ' + code[i])
    #     print('successful_submissions: ' + successful_submissions[i])
    #     print('accuracy: ' + accuracy[i])
    #     print('------------------------------------------------')

    # preparing document object
    document = Document()
    document.add_heading(file_heading, 0)

    for i in range(len(problems)):
        document.add_heading('Problem : ' + str(i + 1) + '\n', level=2)
        document.add_paragraph(
            'heading: ' + str(heading[i]), style='ListBullet'
        )
        document.add_paragraph(
            'code: ' + str(code[i]), style='ListBullet'
        )
        document.add_paragraph(
            'successful submissions: ' + str(successful_submissions[i]), style='ListBullet'
        )
        document.add_paragraph(
            'accuracy: ' + str(accuracy[i]), style='ListBullet'
        )

    document.add_page_break()
    document.save(file_name + '.docx')
    print(file_name + '.docx generated!')


for i in range(len(codechefURLS)):
    PrintProblemsFromURL(codechefURLS[i], file_heading=file_headings[i], file_name=file_names[i])